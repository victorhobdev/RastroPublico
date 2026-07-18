from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).parents[1]
OUT = ROOT / "deliverables"
ASSETS = OUT / "assets"
NAVY = "102A43"
BLUE = "147D92"
TEAL = "2CB1BC"
INK = "243B53"
MUTED = "627D98"
LIGHT = "EAF4F6"
PALE = "F5F8FA"
GOLD = "D9A441"
RED = "B84A4A"


def font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    return ImageFont.truetype(str(next(p for p in candidates if p.exists())), size)


def hexrgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_cell_fill(cell, color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_margins(cell, top=100, start=130, bottom=100, end=130) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar") or OxmlElement("w:tcMar")
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}")) or OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        if node.getparent() is None:
            tc_mar.append(node)
    if tc_mar.getparent() is None:
        tc_pr.append(tc_mar)


def style_run(run, size=10.5, color=INK, bold=False, italic=False, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = hexrgb(color)
    run.bold = bold
    run.italic = italic
    return run


def add_text(doc, text, *, size=10.5, color=INK, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    style_run(p.add_run(text), size, color, bold, italic)
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 16, 2: 13, 3: 11.5}
    before = {1: 14, 2: 10, 3: 7}
    after = {1: 7, 2: 5, 3: 3}
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(before[level])
    p.paragraph_format.space_after = Pt(after[level])
    style_run(p.add_run(text), sizes[level], BLUE if level < 3 else NAVY, True)
    return p


def add_bullet(doc, text, color=INK):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    style_run(p.add_run(text), 10.2, color)
    return p


def add_callout(doc, title, body, color=TEAL):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_fill(cell, LIGHT)
    set_cell_margins(cell, 150, 180, 150, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    style_run(p.add_run(title), 10.5, color, True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.1
    style_run(p2.add_run(body), 10, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_kpi_strip(doc, items):
    table = doc.add_table(rows=1, cols=len(items))
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    width = 6.5 / len(items)
    for cell, (value, label) in zip(table.rows[0].cells, items):
        cell.width = Inches(width)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_fill(cell, NAVY)
        set_cell_margins(cell, 160, 110, 150, 110)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        style_run(p.add_run(value), 17, "FFFFFF", True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        style_run(p2.add_run(label), 8.6, "D9E2EC")


def page_break(doc):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section_geometry(section)
    section.footer.is_linked_to_previous = False
    section.even_page_footer.is_linked_to_previous = False
    section.footer.paragraphs[0].clear()
    add_footer(section)


def make_category_chart(data: dict[str, int], path: Path) -> None:
    labels = [
        ("Licenciamento", data["servico_licenciamento"]),
        ("Computadores e notebooks", data["computador_notebook"]),
        ("Impressoras e scanners", data["impressora_scanner"]),
        ("Rede", data["equipamento_rede"]),
        ("Outsourcing", data["servico_outsourcing"]),
        ("Cloud", data["servico_cloud"]),
        ("Infraestrutura", data["servico_infraestrutura"]),
        ("Monitores", data["monitor"]),
        ("Suporte", data["servico_suporte"]),
        ("Servidores", data["servidor"]),
        ("Desenvolvimento", data["servico_desenvolvimento"]),
    ]
    image = Image.new("RGB", (1500, 760), "white")
    draw = ImageDraw.Draw(image)
    draw.text((45, 28), "Itens tecnológicos por categoria", fill=f"#{NAVY}", font=font(38, True))
    draw.text((45, 78), "Classificação auditada na janela de 12 meses", fill=f"#{MUTED}", font=font(23))
    maximum = max(v for _, v in labels)
    for i, (label, value) in enumerate(labels):
        y = 132 + i * 53
        draw.text((45, y + 7), label, fill=f"#{INK}", font=font(21))
        x0, x1 = 430, 1360
        bar = int((x1 - x0) * value / maximum)
        draw.rounded_rectangle((x0, y, x0 + bar, y + 32), 7, fill=f"#{TEAL if i < 5 else BLUE}")
        draw.text((x0 + bar + 14, y + 3), f"{value:,}".replace(",", "."), fill=f"#{INK}", font=font(21, True))
    image.save(path)


def make_architecture(path: Path) -> None:
    image = Image.new("RGB", (1500, 700), "white")
    draw = ImageDraw.Draw(image)
    draw.text((45, 25), "Da fonte oficial ao indicador", fill=f"#{NAVY}", font=font(38, True))
    boxes = [
        ("Fontes oficiais", "PNCP · Compras.gov\nContratos · contexto"),
        ("Landing", "arquivos imutáveis\nmanifesto + SHA-256"),
        ("Staging", "snapshot Delta\nreconstruível"),
        ("Silver", "tipagem · versão\nqualidade · privacidade"),
        ("Gold", "população elegível\nKPIs + limitações"),
        ("Consumo", "SQL · dashboard\ncase study"),
    ]
    for i, (title, subtitle) in enumerate(boxes):
        x = 40 + i * 242
        fill = NAVY if i in (0, 5) else (BLUE if i in (1, 4) else TEAL)
        draw.rounded_rectangle((x, 190, x + 200, 390), 18, fill=f"#{fill}")
        draw.text((x + 18, 220), title, fill="white", font=font(24, True))
        draw.multiline_text((x + 18, 275), subtitle, fill="#EAF4F6", font=font(19), spacing=9)
        if i < len(boxes) - 1:
            draw.line((x + 202, 290, x + 234, 290), fill=f"#{MUTED}", width=7)
            draw.polygon([(x + 234, 290), (x + 220, 280), (x + 220, 300)], fill=f"#{MUTED}")
    draw.rounded_rectangle((250, 475, 1250, 620), 15, fill=f"#{PALE}", outline=f"#{BLUE}", width=3)
    draw.text((290, 500), "Operação transversal", fill=f"#{BLUE}", font=font(25, True))
    draw.text((290, 548), "run_id · parâmetros · retries · watermark · quarentena · histórico Delta · testes", fill=f"#{INK}", font=font(22))
    image.save(path)


def make_benchmark(path: Path) -> None:
    values = [("Natural + AQE", 3191, TEAL), ("Broadcast hint", 3234, BLUE), ("Sort-merge hint", 6947, RED)]
    image = Image.new("RGB", (1500, 700), "white")
    draw = ImageDraw.Draw(image)
    draw.text((45, 25), "Mesmo workload, estratégias Spark diferentes", fill=f"#{NAVY}", font=font(38, True))
    draw.text((45, 78), "Mediana observada no Databricks serverless · menor é melhor", fill=f"#{MUTED}", font=font(23))
    maxv = max(v for _, v, _ in values)
    for i, (label, value, color) in enumerate(values):
        y = 175 + i * 145
        draw.text((55, y + 25), label, fill=f"#{INK}", font=font(26, True))
        x0, x1 = 430, 1370
        bar = int((x1 - x0) * value / maxv)
        draw.rounded_rectangle((x0, y, x0 + bar, y + 75), 12, fill=f"#{color}")
        draw.text((x0 + bar - 150, y + 20), f"{value / 1000:.2f} s", fill="white", font=font(26, True))
    image.save(path)


def set_section_geometry(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)


def configure(doc: Document) -> None:
    set_section_geometry(doc.sections[0])
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for level, size in ((1, 16), (2, 13), (3, 11.5)):
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = hexrgb(BLUE if level < 3 else NAVY)


def add_footer(section):
    for footer in (section.footer, section.even_page_footer):
        p = footer.paragraphs[0]
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        style_run(p.add_run("RastroPúblico · case study técnico · 18 jul. 2026"), 8.5, MUTED)


def build() -> Path:
    OUT.mkdir(exist_ok=True)
    ASSETS.mkdir(exist_ok=True)
    kpis = json.loads((ROOT / "evidence/data/corrected-kpis.json").read_text(encoding="utf-8"))
    make_category_chart(kpis["itens_por_categoria"], ASSETS / "categorias.png")
    make_architecture(ASSETS / "arquitetura.png")
    make_benchmark(ASSETS / "benchmark.png")

    doc = Document()
    configure(doc)
    for section in doc.sections:
        add_footer(section)

    add_text(doc, "CASE STUDY · ENGENHARIA DE DADOS", size=10, color=TEAL, bold=True, after=34)
    add_text(doc, "RastroPúblico", size=31, color=NAVY, bold=True, after=5)
    add_text(doc, "Compras públicas de tecnologia, com conclusões limitadas pela qualidade real dos dados.", size=16, color=BLUE, after=24)
    add_callout(doc, "O problema", "Transformar milhões de registros oficiais, versões e vínculos imperfeitos em uma base auditável para explorar quem compra de quem, recorrência e evolução contratual — sem classificar fraude ou irregularidade.")
    add_text(doc, "RECORTE AUDITADO", size=9.5, color=MUTED, bold=True, before=15, after=7)
    add_kpi_strip(doc, [("12.252", "compras de tecnologia"), ("29.207", "itens tecnológicos"), ("4.246", "fornecedores distintos"), ("1.537", "contratos vinculados")])
    add_heading(doc, "Decisões que definem o produto", 2)
    add_bullet(doc, "Spark processa arquivos conjuntos, joins, janelas, deduplicação e agregações; Python comum cuida de HTTP.")
    add_bullet(doc, "Landing é imutável; staging é um snapshot Delta reconstruível. A distinção é explícita.")
    add_bullet(doc, "Recorrência exige duas ou mais contratações distintas no mesmo grão.")
    add_bullet(doc, "Totais monetários e preços foram suprimidos após a auditoria revelar semântica não defensável.")
    add_text(doc, "Python · PySpark · Spark SQL · Delta Lake · Databricks Jobs · AI/BI", size=10, color=MUTED, italic=True, before=8)

    page_break(doc)
    add_heading(doc, "1. Arquitetura orientada à evidência", 1)
    add_text(doc, "O pipeline separa evidência bruta, materialização operacional e consumo. Cada camada responde a uma pergunta diferente e pode ser reconstruída sem rebaixar limitações a sucesso.")
    doc.add_picture(str(ASSETS / "arquitetura.png"), width=Inches(6.15))
    add_heading(doc, "O que é realmente incremental", 2)
    add_text(doc, "A coleta, o registro de artefatos e silver.contratacoes usam controle incremental e MERGE. O núcleo Silver e as Gold atuais são reconstruções integrais idempotentes da janela. O projeto não chama overwrite de processamento incremental.")
    add_callout(doc, "Reprocessamento", "Datas, modo e run_id são parâmetros. Watermark só avança após materialização e qualidade; falha preserva as tarefas verdes e permite reparar o trecho afetado.", BLUE)
    add_heading(doc, "Contratos técnicos relevantes", 2)
    add_bullet(doc, "MERGE atômico com uma linha de origem por chave evita ambiguidade e corrida de append.")
    add_bullet(doc, "Hash lógico versionado usa campos ordenados, struct → JSON → SHA-256, exclui metadados técnicos e a privacidade usa whitelist de pessoa jurídica.")

    page_break(doc)
    add_heading(doc, "2. Dados imperfeitos mudam a população — não a narrativa", 1)
    add_text(doc, "Qualidade não é um filtro global. Um item sem unidade pode continuar válido para recorrência, embora seja inelegível para preço. As flags são específicas por indicador e conjuntos vazios recebem estado SEM_DADOS.")
    doc.add_picture(str(ASSETS / "categorias.png"), width=Inches(6.5))
    add_callout(doc, "Leitura permitida", "Licenciamento é a maior categoria classificada na janela. Isso descreve frequência de itens publicados; não mede gasto, risco ou irregularidade.")
    add_heading(doc, "Três resultados defendíveis", 2)
    add_kpi_strip(doc, [("828", "relações recorrentes"), ("20.664", "resultados tecnológicos"), ("1.755", "eventos contratuais")])
    add_text(doc, "Contratos e eventos entram apenas quando ligados a itens tecnológicos. O vínculo contratação–contrato é parcial (cenário C3), então ausência de ligação não significa ausência de contrato.", size=9.7, color=MUTED, italic=True, before=8)

    page_break(doc)
    add_heading(doc, "3. A correção mais importante foi não publicar", 1)
    add_text(doc, "A primeira execução produziu totais financeiramente implausíveis. Em vez de criar um corte arbitrário, a auditoria mediu a distribuição e bloqueou toda conclusão monetária.")
    add_kpi_strip(doc, [("R$ 174,7 tri", "soma bruta nacional"), ("31", "itens acima de R$ 1 tri"), ("R$ 4,23 tri", "maior item tecnológico")])
    add_heading(doc, "Por que o número não virou informação", 2)
    add_bullet(doc, "A fonte mistura semânticas de valor estimado, item, lote e quantidade sem atributos suficientes para resolver todos os casos.")
    add_bullet(doc, "Equivalência entre duas consultas confirma que ambas somam o mesmo dado; não confirma que o campo significa o que se deseja publicar.")
    add_bullet(doc, "Não existe limiar universal defensável para remover outliers; logo, ranking, total financeiro e comparação de preço permanecem desabilitados.")
    add_callout(doc, "Decisão de governança", "O produto publica contagens, recorrência e cobertura. Preço e valor só voltarão após um gate mensurável de unidade, escopo, atributos e revisão amostral.", GOLD)
    add_heading(doc, "O que essa decisão demonstra", 2)
    add_text(doc, "Investigar semântica antes de otimizar ou apresentar; registrar a limitação; e preferir uma ausência explicada a uma conclusão precisa na aparência e errada no significado.")

    page_break(doc)
    add_heading(doc, "4. Spark foi medido, não decorativo", 1)
    add_text(doc, "O benchmark compara estratégias Spark no mesmo ambiente e snapshot. Não é um duelo Pandas × Spark e não generaliza o resultado para outros workloads.")
    doc.add_picture(str(ASSETS / "benchmark.png"), width=Inches(6.5))
    add_kpi_strip(doc, [("4.791.466", "linhas lidas"), ("39", "arquivos"), ("270,53 MB", "shuffle observado"), ("0", "spill")])
    add_heading(doc, "Conclusão técnica", 2)
    add_text(doc, "Manter a estratégia natural com AQE. O hint de broadcast foi praticamente equivalente; forçar sort-merge mais que dobrou a mediana. O experimento registrou warm-up, ordem alternada, cache de resultado desabilitado, checksum e planos inicial/executado disponíveis conforme a interface.")
    add_callout(doc, "Limite da evidência", "O repositório contém Job, run, Query History e output do notebook. O JSON detalhado do Query Profile não foi exportado; portanto, essa evidência não é alegada como presente.", BLUE)

    page_break(doc)
    add_heading(doc, "5. O que um revisor consegue verificar", 1)
    add_text(doc, "O repositório foi reorganizado para que alegações importantes apontem para código, teste ou evidência exportada — e para que relatórios históricos não sejam confundidos com a baseline corrigida.")
    rows = [
        ("KPIs corrigidos", "evidence/data/corrected-kpis.json", "Recalcula população tecnológica e recorrência"),
        ("Semântica monetária", "evidence/data/value-semantics-summary.json", "Explica por que valores não são publicados"),
        ("Jobs e benchmark", "databricks.yml + evidence/databricks/", "Parâmetros, run e métricas sanitizadas"),
        ("Contratos executáveis", "tests/", "109 testes aprovados na auditoria modular"),
        ("Ambiente limpo", ".github/workflows/ci.yml", "uv.lock, lint, suíte e cobertura mínima"),
    ]
    for evidence, location, proof in rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        style_run(p.add_run(f"{evidence}  "), 9.7, NAVY, True)
        style_run(p.add_run(location), 9.2, BLUE)
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Inches(0.2)
        p2.paragraph_format.space_after = Pt(5)
        style_run(p2.add_run(proof), 9.3, MUTED)
    add_heading(doc, "Limitações que permanecem visíveis", 2)
    add_bullet(doc, "As tabelas Delta corrigidas precisam ser rematerializadas quando houver compute disponível.")
    add_bullet(doc, "O runtime local (Spark 4.2) difere do serverless observado (Spark 4.1), embora o código use APIs compatíveis.")
    add_bullet(doc, "Cobertura de modalidades e vínculo contratual são parciais; o produto não representa todo o universo PNCP.")
    add_bullet(doc, "A tabela órgão–fornecedor é uma lista de arestas, não uma análise de grafos.")
    add_callout(doc, "Síntese", "A principal evidência profissional não é o volume ou a quantidade de tecnologias. É a capacidade de encontrar uma conclusão errada, rastrear sua causa, corrigir a população e reduzir a alegação ao que os dados sustentam.")

    output = OUT / "RastroPublico-case-study.docx"
    doc.save(output)
    return output


if __name__ == "__main__":
    print(build())
